#backend/test/test_extract_regulations.pyimport pytest
from pathlib import Path
import pytest
from backend.scripts.extract_regulations import (
    extract_text_from_pdf,
    extract_text_from_docx,
    extract_text,
)

def test_extract_text_file_not_found():
    """Checks that FileNotFoundError is raised if the file does not exist."""
    with pytest.raises(FileNotFoundError):
        extract_text(Path("not_existing_file.pdf"))

def test_extract_text_invalid_extension(tmp_path):
    """Checks that ValueError is raised if the file extension is not supported."""
    file_path = tmp_path / "invalid.txt"
    file_path.write_text("dummy content")
    with pytest.raises(ValueError):
        extract_text(file_path)

def test_extract_docx_basic(tmp_path):
    """Checks that text is correctly extracted from a DOCX file."""
    from docx import Document

    file_path = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("Hello Regulations")
    doc.save(file_path)

    result = extract_text_from_docx(file_path)
    assert "Hello Regulations" in result

def test_extract_pdf_basic(tmp_path):
    """Checks that text is correctly extracted from a PDF file."""
    import fitz

    file_path = tmp_path / "sample.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Hello PDF Regulations")
    doc.save(file_path)

    result = extract_text_from_pdf(file_path)
    assert "Hello PDF Regulations" in result


def test_extract_text_docx_valid(tmp_path):
    """Checks that extract_text() works correctly with a valid DOCX file."""
    from docx import Document

    file_path = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("Regulation Check")
    doc.save(file_path)

    result = extract_text(file_path)
    assert "Regulation Check" in result